import os, requests, json, re, logging
from dotenv import load_dotenv
from .utils import *
from .db_utils import *
from logging.handlers import RotatingFileHandler
from docx import Document
from docx.shared import Pt

# ------------------------------------------------------------------
# Environment
# ------------------------------------------------------------------
load_dotenv()
OPENROUTER_API_KEY = os.environ["OPENROUTER_API_KEY"]

# ------------------------------------------------------------------
# Logging Configuration
# ------------------------------------------------------------------
LOG_LEVEL = os.getenv("BOOKGEN_LOG_LEVEL", "INFO").upper()
LOG_DIR = os.getenv("BOOKGEN_LOG_DIR", "logs")
LOG_FILE = os.path.join(LOG_DIR, "bookgen.log")

os.makedirs(LOG_DIR, exist_ok=True)

logger = logging.getLogger("BookGen")
logger.setLevel(LOG_LEVEL)

formatter = logging.Formatter(
    "%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)

# ---------------- Console Handler ----------------
console_handler = logging.StreamHandler()
console_handler.setLevel(LOG_LEVEL)
console_handler.setFormatter(formatter)

# ---------------- File Handler ----------------
file_handler = RotatingFileHandler(
    LOG_FILE,
    maxBytes=10 * 1024 * 1024,  # 10 MB
    backupCount=5,
    encoding="utf-8"
)
file_handler.setLevel(logging.DEBUG)  # capture everything in file
file_handler.setFormatter(formatter)

# ---------------- Attach Handlers ----------------
logger.handlers.clear()      # avoid duplicate logs
logger.addHandler(console_handler)
logger.addHandler(file_handler)
logger.propagate = False

# ------------------------------------------------------------------
# BookGen Class
# ------------------------------------------------------------------
class BookGen:
    def __init__(self, OPENROUTER_API_KEY=OPENROUTER_API_KEY, model="deepseek/deepseek-r1-0528:free"):
        self.model = model
        self._api_key = OPENROUTER_API_KEY
        self.outline_template = None
        self.content_template = None
        self.summarize_template = None
        self.conn = get_db_connection()
        logger.info("BookGen initialized")
        logger.info(f"Using model: {self.model}")

    # ------------------------------------------------------------------
    # Prompt Templates
    # ------------------------------------------------------------------
    def get_outline_template(self):
        logger.debug("Loading OUTLINE system prompt")
        return [{"role": "system", "content": OUTLINE_SYS_PROMPT}]

    def get_content_template(self):
        logger.debug("Loading CONTENT system prompt")
        return [{"role": "system", "content": CONTENT_SYS_PROMPT}]
    
    def get_summarize_template(self):
        logger.debug("Loading SUMMARIZE system prompt")
        return [{"role": "system", "content": SUMMARIZE_SYS_PROMPT}]

    # ------------------------------------------------------------------
    # LLM JSON Parsing
    # ------------------------------------------------------------------
    def parse_llm_json(self, response_text: str) -> dict:
        logger.debug("Parsing JSON from LLM response")
        match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if not match:
            logger.error("No JSON detected in LLM response")
            logger.debug(f"RAW RESPONSE:\n{response_text}")
            raise ValueError("No JSON found in response")
        try:
            parsed = json.loads(match.group())
            logger.debug("JSON parsed successfully")
            return parsed
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {e}")
            logger.debug(f"RAW RESPONSE:\n{response_text}")
            raise ValueError(f"Invalid JSON format: {e}")

    # ------------------------------------------------------------------
    # LLM Call
    # ------------------------------------------------------------------
    def call_model(self, message):
        logger.debug(
            "Users Prompt:\n%s",
            message[-1]["content"] if len(message[-1]["content"]) < 8000 else message[-1]["content"][:8000] + "\n--- TRUNCATED ---"
        )
        logger.info("Calling OpenRouter LLM API")
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {self._api_key}",
                "Content-Type": "application/json",
            },
            data=json.dumps({
                "model": self.model,
                "messages": message,
            })
        )
        response_data = response.json()
        if "choices" not in response_data:
            logger.error("Invalid OpenRouter response structure")
            logger.error(response_data)
            raise RuntimeError("Invalid response from OpenRouter")
        llm_output = response_data["choices"][0]["message"]["content"]
        # Log full LLM response (truncate only if extremely large)
        logger.debug(
            "LLM RESPONSE:\n%s",
            llm_output if len(llm_output) < 8000 else llm_output[:8000] + "\n--- TRUNCATED ---"
        )
        return llm_output

    # ------------------------------------------------------------------
    # Outline Generation
    # ------------------------------------------------------------------
    def generate_outline(self, title, notes):
        logger.info(f"Generating outline for book: {title}")
        prompt = OUTLINE_PROMPT.replace("__BOOK_TITLE__", title)\
                               .replace("__EDITORIAL_NOTES__", notes)
        self.outline_template = self.get_outline_template() if not self.outline_template else self.outline_template
        for attempt in range(1, 4):
            try:
                logger.info(f"Outline attempt {attempt}")
                self.outline_template.append({"role": "user", "content": prompt})
                outline = self.call_model(self.outline_template)
                parsed_outline = self.parse_llm_json(outline)
                logger.info("Outline generated successfully")
                return json.dumps(parsed_outline, indent=2)
            except ValueError as e:
                logger.warning(f"Outline parsing failed (attempt {attempt}): {e}")
        logger.critical("Outline generation failed after retries")
        raise RuntimeError("Failed to generate a valid outline")

    # ------------------------------------------------------------------
    # Save Book & Outline
    # ------------------------------------------------------------------
    def save_book_and_outline(self, title, notes, outline_json):
        logger.info(f"Saving book: {title}")
        if get_book(self.conn, title):
            logger.warning(f"Book '{title}' exists. Overwriting.")
            delete_book(self.conn, title)
        book_id = add_book(self.conn, title, notes)
        logger.info(f"Book saved with ID: {book_id}")
        outline_data = json.loads(outline_json)
        for chapter in outline_data.get("outline", []):
            add_heading(
                self.conn,
                book_id=book_id,
                heading_title=chapter["chapter_title"],
                heading_number=chapter["chapter_number"],
                sub_heading="\n".join(chapter["sections"]),
                description=chapter["chapter_description"]
            )
            logger.debug(f"Chapter saved: {chapter['chapter_title']}")
        logger.info("All chapters saved successfully")
        return book_id

    # ------------------------------------------------------------------
    # Content & Summary Generation
    # ------------------------------------------------------------------
    def generate_heading_content(self, book_title, heading_notes=None, starting_heading_number=1):
        logger.info(f"Generating content for book: {book_title}")
        book = get_book(self.conn, book_title)
        if not book:
            logger.error(f"Book not found: {book_title}")
            raise ValueError(f"No book found with title {book_title}")
        headings = get_headings_by_book(self.conn, book["id"])
        previous_summary = {}
        self.content_template = self.get_content_template() if not self.content_template else self.content_template
        self.summarize_template = self.get_summarize_template() if not self.summarize_template else self.summarize_template
        for i in range(starting_heading_number - 1):
            previous_summary[headings[i]["heading_title"]] = headings[i]["summary"]
        for i in range(starting_heading_number - 1, len(headings)):
            heading = headings[i]
            logger.info(f"Processing heading: {heading['heading_title']}")
            content_prompt = CONTENT_PROMPT.replace(
                "__BOOK_TITLE__", book["title"]
            ).replace(
                "__HEADING_TITLE__", heading["heading_title"]
            ).replace(
                "__SUB_HEADINGS__", heading["sub_heading"]
            ).replace(
                "__PREVIOUS_HEADINGS_SUMMARY_DICT__", json.dumps(previous_summary)
            )
            if heading_notes and heading["heading_title"] in heading_notes:
                content_prompt = content_prompt.replace(
                    "__EDITORIAL_NOTES__", heading_notes[heading["heading_title"]]
                )
            # Generate content
            for attempt in range(1, 4):
                try:
                    logger.debug(f"Content generation attempt {attempt}")
                    self.content_template.append({"role": "user", "content": content_prompt})
                    content = self.call_model(self.content_template)
                    break
                except Exception as e:
                    logger.warning(f"Content generation failed (attempt {attempt}): {e}")
            # Generate summary
            summary_prompt = SUMMARIZE_PROMPT.replace(
                "__BOOK_TITLE__", book["title"]
            ).replace(
                "__HEADING_TITLE__", heading["heading_title"]
            ).replace(
                "__INPUT_TEXT__", content
            )
            if heading_notes and heading["heading_title"] in heading_notes:
                summary_prompt = summary_prompt.replace(
                    "__EDITORIAL_NOTES__", heading_notes[heading["heading_title"]]
                )
            for attempt in range(1, 4):
                try:
                    logger.debug(f"Summary generation attempt {attempt}")
                    self.summarize_template.append({"role": "user", "content": summary_prompt})
                    summary = self.call_model(self.summarize_template)
                    break
                except Exception as e:
                    logger.warning(f"Summary generation failed (attempt {attempt}): {e}")
            update_heading(
                self.conn,
                book_id=book["id"],
                heading_title=heading["heading_title"],
                summary=summary,
                content=content
            )
            previous_summary[heading["heading_title"]] = summary
            logger.info(f"Completed heading: {heading['heading_title']}")

    # ------------------------------------------------------------------
    # Retrieve Book
    # ------------------------------------------------------------------
    def get_book_and_outline(self, book_title):
        logger.info(f"Fetching book: {book_title}")
        book = get_book(self.conn, book_title)
        if not book:
            logger.error(f"Book not found: {book_title}")
            raise ValueError(f"No book found with title {book_title}")
        headings = get_headings_by_book(self.conn, book["id"])
        outline = {
            "book_title": book["title"],
            "outline": []
        }
        for heading in headings:
            outline["outline"].append({
                "chapter_number": heading["heading_number"],
                "chapter_title": heading["heading_title"],
                "chapter_description": heading["description"],
                "sections": heading["sub_heading"],
                "summary": heading["summary"],
                "content": heading["content"]
            })
        logger.info("Book outline retrieved successfully")
        return json.dumps(outline, indent=2)
    
    def book_gen(self, title, output_path):
        doc = Document()
        # Basic document styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        doc.add_heading(title.title(), level=0)
        logger.info(f"Fetching book: {title}")
        book = get_book(self.conn, title)
        headings = get_headings_by_book(self.conn, book["id"])
        for heading in headings:
            doc.add_heading(f"{heading['heading_number']}. {heading['heading_title']}", level=1)
            doc.add_paragraph(heading["content"])
        doc.save(output_path)

# ------------------------------------------------------------------
# Entry Point
# ------------------------------------------------------------------
if __name__ == "__main__":
    book_gen = BookGen()
    sample_title = "The Future of Artificial Intelligence"
    # sample_notes = (
    #     "Explore the advancements in AI technology, its applications across various industries, "
    #     "ethical considerations, and potential future developments."
    # )
    # outline = book_gen.generate_outline(sample_title, sample_notes)
    # book_gen.save_book_and_outline(sample_title, sample_notes, outline)
    # print(book_gen.get_book_and_outline(sample_title))
    # start = int(input("Enter starting chapter number for content generation (1 for beginning): ") or "1")
    # book_gen.generate_heading_content(
    #     sample_title,
    #     starting_heading_number=start,
    # )
    # print(book_gen.get_book_and_outline(sample_title))
    book_gen.book_gen(sample_title, "output.docx")
